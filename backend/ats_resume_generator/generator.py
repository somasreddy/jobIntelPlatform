import io
import logging

logger = logging.getLogger(__name__)

_SYSTEM_PROMPT = """You are a World-Class ATS Resume Optimization Specialist.
Transform the candidate's profile into a high-impact, professionally tailored resume for the specific job description.
This must work for ANY role — engineering, product, design, sales, marketing, finance, or any other.

REQUIRED OUTPUT (JSON):
1. summary: A 3-line powerful professional summary using JD language. Formula: "Years experience + Key Skills + Major Achievement".
2. bullets: Exactly 5 quantified experience bullets using Problem-Action-Result (PAR) framework.
   - Start each with a strong action verb (Spearheaded, Engineered, Orchestrated, Drove, Reduced).
   - MUST include metrics (%, $, time saved, scale, frequency).
   - Use the exact terminology from the job description.
3. skills_grouped: Dictionary of 3-4 relevant categories with skills from both the profile and JD.

Rules:
- Mirror the JD's exact technical/domain terminology.
- Prioritize keywords that appear multiple times in the JD.
- Output ONLY the JSON object, no markdown, no preamble.
"""


async def _generate_content(
    profile: dict, job_description: str, tech_stack: list
) -> dict:
    try:
        from core.llm import smart_chat
        user_prompt = (
            f"Candidate Role: {profile.get('current_role', 'Professional')}\n"
            f"Experience: {profile.get('experience_years', 5)} years\n"
            f"Profile Skills: {', '.join((profile.get('skills') or []))}\n"
            f"Profile Frameworks: {', '.join((profile.get('frameworks') or []))}\n"
            f"Target Tech Stack / Keywords: {', '.join(tech_stack[:15])}\n\n"
            f"Full Job Description:\n{job_description[:3000]}\n\n"
            "Generate 'summary', 'bullets' (PAR format), and 'skills_grouped' as a JSON object."
        )
        raw = await smart_chat(_SYSTEM_PROMPT, user_prompt, temperature=0.5, task_type="resume", max_tokens=2000)
        import json
        clean = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        data = json.loads(clean)
        if isinstance(data, dict):
            return data
    except Exception as e:
        logger.warning(f"LLM content generation failed: {e}")

    # Fallback
    return {
        "summary": f"Highly skilled {profile.get('current_role', 'QA')} professional with {profile.get('experience_years', 5)}+ years experience, specializing in {tech_stack[0] if tech_stack else 'automation'} and delivery excellence.",
        "bullets": [
            f"Optimized {tech_stack[0] if tech_stack else 'test'} execution pipelines by 40% through strategic implementation of parallel processing and containerization.",
            "Led a cross-functional quality initiative that reduced production defect leakage by 25% across 3 critical microservices.",
            "Architected an end-to-end automation suite from scratch, achieving 85% regression coverage within the first 3 months.",
            "Spearheaded the integration of contract testing into CI/CD, preventing integration issues and saving 10+ engineering hours per week.",
            "Mentored 5+ junior engineers in modern automation best practices, increasing team velocity by 30%."
        ],
        "skills_grouped": {
            "Automation": tech_stack[:4],
            "Tools": ["Docker", "Kubernetes", "Git"],
            "Languages": ["Java", "Python", "TypeScript"]
        }
    }


def _build_pdf(profile: dict, job: dict, bullets: list, summary: str = "", skills_grouped: dict = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib import colors

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
            leftMargin=1.8*cm, rightMargin=1.8*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        story = []

        # Premium Styles
        name_style = ParagraphStyle("Name", parent=styles["Heading1"], fontSize=22, spaceAfter=2,
            textColor=colors.HexColor("#1e293b"), fontName="Helvetica-Bold")
        sub_header_style = ParagraphStyle("SubHeader", parent=styles["Normal"], fontSize=10, 
            textColor=colors.HexColor("#64748b"), spaceAfter=10)
        section_style = ParagraphStyle("Section", parent=styles["Heading2"], fontSize=11,
            spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#4f46e5"), fontName="Helvetica-Bold",
            borderPadding=(0, 0, 2, 0))
        body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9.5, leading=14, spaceAfter=6, textColor=colors.HexColor("#334155"))
        bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=9.5, leading=14,
            leftIndent=15, spaceAfter=4, textColor=colors.HexColor("#334155"))
        skill_cat_style = ParagraphStyle("SkillCat", parent=styles["Normal"], fontSize=9, fontName="Helvetica-Bold", textColor=colors.HexColor("#1e293b"))

        name = profile.get("name") or "Candidate"
        role = profile.get("current_role", "QA Engineer")
        loc = profile.get("current_location", "")
        exp = profile.get("experience_years", 0)

        # Header
        story.append(Paragraph(name, name_style))
        story.append(Paragraph(f"{role}  |  {loc}  |  {exp}+ Years Experience", sub_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=10))

        # Professional Summary
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        story.append(Paragraph(summary or f"Results-driven {role} with {exp}+ years experience.", body_style))

        # Technical Skills (Categorized)
        story.append(Paragraph("CORE COMPETENCIES", section_style))
        if skills_grouped:
            for cat, items in skills_grouped.items():
                story.append(Paragraph(f"<b>{cat}:</b> {', '.join(items)}", body_style))
        else:
            skills = profile.get("skills") or []
            story.append(Paragraph(", ".join(skills[:15]), body_style))

        # Experience
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_style))
        story.append(Paragraph(f"<b>{role}</b>  |  {job.get('organization', 'Current Company')}", body_style))
        for b in bullets:
            story.append(Paragraph(f"• {b}", bullet_style))

        doc.build(story)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"PDF build error: {e}")
        return b""


def _build_docx(profile: dict, job: dict, bullets: list, summary: str = "", skills_grouped: dict = None) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        name = profile.get("name") or "Candidate"
        role = profile.get("current_role", "QA Engineer")
        loc = profile.get("current_location", "")
        exp = profile.get("experience_years", 0)

        h = doc.add_heading(name, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"{role}  |  {loc}  |  {exp}+ Years Experience").font.size = Pt(10)

        def add_section(title: str):
            p = doc.add_heading(title, level=1)
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

        add_section("Professional Summary")
        doc.add_paragraph(summary or f"Experienced {role} with {exp}+ years experience.")

        add_section("Core Competencies")
        if skills_grouped:
            for cat, items in skills_grouped.items():
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{cat}: ").bold = True
                p.add_run(", ".join(items))
        else:
            skills = profile.get("skills") or []
            doc.add_paragraph(", ".join(skills[:15]))

        add_section("Professional Experience")
        p = doc.add_paragraph()
        p.add_run(f"{role}  |  {job.get('organization', 'Current Company')}").bold = True
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"DOCX build error: {e}")
        return b""


def _build_master_pdf(
    profile: dict,
    summary: str,
    experience_blocks: list[dict],
    education_blocks: list[dict] = None,
    achievements: list[str] = None,
    contact: dict = None,
) -> bytes:
    """Build a comprehensive master ATS resume PDF from full profile data."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, HRFlowable
        from reportlab.lib import colors

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
            leftMargin=1.8*cm, rightMargin=1.8*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        story = []

        accent = colors.HexColor("#4f46e5")
        dark   = colors.HexColor("#1e293b")
        muted  = colors.HexColor("#64748b")
        body_dark = colors.HexColor("#334155")
        divider = colors.HexColor("#e2e8f0")

        name_style     = ParagraphStyle("Name",     parent=styles["Heading1"], fontSize=24, spaceAfter=2, textColor=dark, fontName="Helvetica-Bold")
        sub_style      = ParagraphStyle("Sub",      parent=styles["Normal"],   fontSize=9.5, textColor=muted, spaceAfter=3)
        contact_style  = ParagraphStyle("Cont",     parent=styles["Normal"],   fontSize=9,   textColor=muted, spaceAfter=10)
        section_style  = ParagraphStyle("Section",  parent=styles["Heading2"], fontSize=11, spaceBefore=12, spaceAfter=4, textColor=accent, fontName="Helvetica-Bold")
        body_style     = ParagraphStyle("Body",     parent=styles["Normal"],   fontSize=9.5, leading=14, spaceAfter=5, textColor=body_dark)
        bullet_style   = ParagraphStyle("Bullet",   parent=styles["Normal"],   fontSize=9.5, leading=14, leftIndent=12, spaceAfter=3, textColor=body_dark)
        exp_title_style= ParagraphStyle("ExpTitle", parent=styles["Normal"],   fontSize=10,  fontName="Helvetica-Bold", textColor=dark, spaceAfter=1, spaceBefore=7)
        exp_meta_style = ParagraphStyle("ExpMeta",  parent=styles["Normal"],   fontSize=9,   textColor=muted, spaceAfter=3)

        name      = profile.get("name") or "Candidate"
        role      = profile.get("current_role") or ""
        loc       = profile.get("current_location") or ""
        exp       = profile.get("experience_years") or 0
        work_mode = profile.get("work_mode") or ""
        certs     = profile.get("certifications") or []

        # ── Header ──────────────────────────────────────────────────────────
        story.append(Paragraph(name, name_style))
        meta_parts = [p for p in [role, loc, f"{exp}+ yrs" if exp else "", work_mode] if p]
        story.append(Paragraph("  ·  ".join(meta_parts), sub_style))

        # Contact line
        ct = contact or {}
        contact_parts = [p for p in [ct.get("email"), ct.get("phone"), ct.get("linkedin"), ct.get("github")] if p]
        if contact_parts:
            story.append(Paragraph("  |  ".join(contact_parts), contact_style))
        else:
            story.append(Paragraph("", contact_style))

        story.append(HRFlowable(width="100%", thickness=1.5, color=accent, spaceAfter=8))

        # ── Professional Summary ─────────────────────────────────────────────
        if summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            story.append(Paragraph(summary, body_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=divider, spaceAfter=4))

        # ── Technical Skills ─────────────────────────────────────────────────
        skills    = profile.get("skills") or []
        frameworks= profile.get("frameworks") or []
        languages = profile.get("languages") or []
        cicd      = profile.get("cicd_tools") or []
        ai_tools  = profile.get("ai_tools") or []

        skill_rows = []
        if skills:     skill_rows.append(("Core Skills",    ", ".join(skills)))
        if frameworks: skill_rows.append(("Frameworks",     ", ".join(frameworks)))
        if languages:  skill_rows.append(("Languages",      ", ".join(languages)))
        if cicd:       skill_rows.append(("CI/CD & DevOps", ", ".join(cicd)))
        if ai_tools:   skill_rows.append(("AI Tools",       ", ".join(ai_tools)))

        if skill_rows:
            story.append(Paragraph("TECHNICAL SKILLS", section_style))
            for cat, items in skill_rows:
                story.append(Paragraph(f"<b>{cat}:</b>  {items}", body_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=divider, spaceAfter=4))

        # ── Professional Experience ──────────────────────────────────────────
        if experience_blocks:
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_style))
            for block in experience_blocks:
                title_line = block.get("title", "")
                meta_line  = "  |  ".join(p for p in [block.get("company", ""), block.get("duration", "")] if p)
                if title_line:
                    story.append(Paragraph(title_line, exp_title_style))
                if meta_line:
                    story.append(Paragraph(meta_line, exp_meta_style))
                for bullet in (block.get("bullets") or []):
                    story.append(Paragraph(f"• {bullet}", bullet_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=divider, spaceAfter=4))

        # ── Education ────────────────────────────────────────────────────────
        edu = education_blocks or []
        if edu:
            story.append(Paragraph("EDUCATION", section_style))
            for ed in edu:
                degree = ed.get("degree", "")
                inst   = ed.get("institution", "")
                year   = ed.get("year", "")
                edloc  = ed.get("location", "")
                if degree:
                    story.append(Paragraph(degree, exp_title_style))
                meta = "  |  ".join(p for p in [inst, edloc, year] if p)
                if meta:
                    story.append(Paragraph(meta, exp_meta_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=divider, spaceAfter=4))

        # ── Certifications ───────────────────────────────────────────────────
        if certs:
            story.append(Paragraph("CERTIFICATIONS", section_style))
            for cert in certs:
                story.append(Paragraph(f"• {cert}", bullet_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=divider, spaceAfter=4))

        # ── Achievements ─────────────────────────────────────────────────────
        ach = achievements or []
        if ach:
            story.append(Paragraph("ACHIEVEMENTS & AWARDS", section_style))
            for item in ach:
                story.append(Paragraph(f"• {item}", bullet_style))

        doc.build(story)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"Master PDF build error: {e}")
        return b""


def _build_master_docx(
    profile: dict,
    summary: str,
    experience_blocks: list[dict],
    education_blocks: list[dict] = None,
    achievements: list[str] = None,
    contact: dict = None,
) -> bytes:
    """Build a comprehensive master ATS resume DOCX from full profile data."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        accent_rgb = RGBColor(0x4F, 0x46, 0xE5)
        dark_rgb   = RGBColor(0x1E, 0x29, 0x3B)
        muted_rgb  = RGBColor(0x64, 0x74, 0x8B)

        name      = profile.get("name") or "Candidate"
        role      = profile.get("current_role") or ""
        loc       = profile.get("current_location") or ""
        exp       = profile.get("experience_years") or 0
        work_mode = profile.get("work_mode") or ""
        certs     = profile.get("certifications") or []
        ct        = contact or {}

        # Name
        h = doc.add_heading(name, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.runs[0].font.size = Pt(22)
        h.runs[0].font.color.rgb = dark_rgb

        # Sub-header
        meta_parts = [p for p in [role, loc, f"{exp}+ yrs" if exp else "", work_mode] if p]
        sub = doc.add_paragraph("  ·  ".join(meta_parts))
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = muted_rgb

        # Contact line
        contact_parts = [p for p in [ct.get("email"), ct.get("phone"), ct.get("linkedin"), ct.get("github")] if p]
        if contact_parts:
            cp = doc.add_paragraph("  |  ".join(contact_parts))
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = muted_rgb

        def add_section_heading(title: str):
            p = doc.add_heading(title, level=1)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = accent_rgb
            p.runs[0].font.bold = True

        def add_bold_sub(text: str, size: float = 10.5):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(size)
            r.font.color.rgb = dark_rgb

        def add_meta_line(text: str):
            p = doc.add_paragraph(text)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = muted_rgb

        # Summary
        if summary:
            add_section_heading("Professional Summary")
            doc.add_paragraph(summary)

        # Skills
        skills    = profile.get("skills") or []
        frameworks= profile.get("frameworks") or []
        languages = profile.get("languages") or []
        cicd      = profile.get("cicd_tools") or []
        ai_tools  = profile.get("ai_tools") or []

        skill_rows = []
        if skills:     skill_rows.append(("Core Skills",    ", ".join(skills)))
        if frameworks: skill_rows.append(("Frameworks",     ", ".join(frameworks)))
        if languages:  skill_rows.append(("Languages",      ", ".join(languages)))
        if cicd:       skill_rows.append(("CI/CD & DevOps", ", ".join(cicd)))
        if ai_tools:   skill_rows.append(("AI Tools",       ", ".join(ai_tools)))

        if skill_rows:
            add_section_heading("Technical Skills")
            for cat, items in skill_rows:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"{cat}: ")
                r.bold = True
                r.font.color.rgb = dark_rgb
                p.add_run(items)

        # Experience
        if experience_blocks:
            add_section_heading("Professional Experience")
            for block in experience_blocks:
                title_line = block.get("title", "")
                company    = block.get("company", "")
                duration   = block.get("duration", "")
                if title_line:
                    add_bold_sub(title_line)
                if company or duration:
                    add_meta_line("  |  ".join(x for x in [company, duration] if x))
                for bullet in (block.get("bullets") or []):
                    doc.add_paragraph(bullet, style="List Bullet")

        # Education
        edu = education_blocks or []
        if edu:
            add_section_heading("Education")
            for ed in edu:
                degree = ed.get("degree", "")
                inst   = ed.get("institution", "")
                year   = ed.get("year", "")
                edloc  = ed.get("location", "")
                if degree:
                    add_bold_sub(degree, size=10)
                meta = "  |  ".join(p for p in [inst, edloc, year] if p)
                if meta:
                    add_meta_line(meta)

        # Certifications
        if certs:
            add_section_heading("Certifications")
            for cert in certs:
                doc.add_paragraph(cert, style="List Bullet")

        # Achievements
        ach = achievements or []
        if ach:
            add_section_heading("Achievements & Awards")
            for item in ach:
                doc.add_paragraph(item, style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"Master DOCX build error: {e}")
        return b""


_MASTER_SUMMARY_PROMPT = """You are a professional resume writer.
Given a candidate profile, write a 3-4 sentence professional summary for their master resume.
The summary should highlight years of experience, core expertise, key tools/technologies, and value proposition.
Write in first-person implied style (no "I"). Be specific and impactful.
IMPORTANT: Return ONLY the plain summary paragraph text. No JSON. No code blocks. No labels. No quotes."""

_RESUME_PARSE_PROMPT = """You are an expert resume parser. Extract all structured sections from the provided resume text.
Return ONLY a valid JSON object with exactly this shape (use empty arrays if a section is not found):
{
  "experience": [
    {
      "title": "Job Title",
      "company": "Company Name",
      "duration": "Mon YYYY – Mon YYYY",
      "bullets": ["Specific accomplishment or responsibility 1", "Accomplishment 2", "Accomplishment 3"]
    }
  ],
  "education": [
    {
      "degree": "Degree / Qualification",
      "institution": "Institution Name",
      "year": "YYYY",
      "location": "City, State"
    }
  ],
  "achievements": ["Award or recognition 1", "Achievement 2"],
  "contact": {
    "email": "",
    "phone": "",
    "linkedin": "",
    "github": ""
  }
}
Each experience entry should have 3-5 bullets capturing projects and accomplishments.
Return valid JSON only — no markdown, no code blocks, no preamble."""


def _extract_json_safely(raw: str) -> dict | list | None:
    """Extract JSON from LLM output that may be wrapped in markdown code blocks."""
    import json as _json, re
    text = raw.strip()
    # Strip code fences
    text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\s*```$', '', text, flags=re.MULTILINE)
    text = text.strip()
    try:
        return _json.loads(text)
    except Exception:
        # Try to find the first { or [ and parse from there
        for start_char, end_char in [('{', '}'), ('[', ']')]:
            idx = text.find(start_char)
            if idx != -1:
                try:
                    return _json.loads(text[idx:])
                except Exception:
                    pass
    return None


def _clean_summary(raw: str) -> str:
    """Ensure summary is plain text — strip JSON wrappers if LLM misbehaved."""
    import json as _json, re
    text = raw.strip().strip('"').strip("'")
    # If it looks like JSON, try to extract the summary field
    if text.startswith('{') or text.startswith('```'):
        parsed = _extract_json_safely(text)
        if isinstance(parsed, dict):
            for key in ('summary', 'professional_summary', 'text', 'content'):
                if parsed.get(key) and isinstance(parsed[key], str):
                    return parsed[key].strip()
        # Strip any remaining code-block syntax and return as-is
        text = re.sub(r'```(?:json)?\s*|\s*```', '', text).strip()
    return text


async def generate_master_resume(profile: dict) -> dict:
    """Generate a comprehensive master ATS resume PDF + DOCX from profile + resumeText."""
    import base64, asyncio

    resume_text = profile.get("base_resume_text") or profile.get("resumeText") or profile.get("resume_text") or ""

    from core.llm import smart_chat

    # ── 1. Generate professional summary via LLM ──────────────────────────────
    summary = ""
    try:
        user_msg = (
            f"Name: {profile.get('name', '')}\n"
            f"Role: {profile.get('current_role', '')}\n"
            f"Experience: {profile.get('experience_years', 0)} years\n"
            f"Location: {profile.get('current_location', '')}\n"
            f"Skills: {', '.join((profile.get('skills') or [])[:12])}\n"
            f"Frameworks: {', '.join((profile.get('frameworks') or [])[:8])}\n"
            f"Languages: {', '.join((profile.get('languages') or [])[:6])}\n"
            f"CI/CD Tools: {', '.join((profile.get('cicd_tools') or [])[:6])}\n"
            + (f"Resume excerpt:\n{resume_text[:1200]}" if resume_text else "")
        )
        raw_summary = await smart_chat(
            _MASTER_SUMMARY_PROMPT, user_msg,
            max_tokens=350, temperature=0.4,
            task_type="resume_parse", cache_ttl=0,
        )
        summary = _clean_summary(raw_summary)
    except Exception as e:
        logger.warning(f"Master summary LLM failed: {e}")

    if not summary:
        role = profile.get("current_role", "Professional")
        exp = profile.get("experience_years", 0)
        skills_preview = ", ".join((profile.get("skills") or [])[:5])
        summary = (
            f"Results-driven {role} with {exp}+ years of experience"
            + (f" in {skills_preview}" if skills_preview else "")
            + ". Proven track record of delivering high-quality solutions, driving automation, and enabling measurable business impact across cross-functional teams."
        )

    # ── 2. Parse all sections from resumeText ────────────────────────────────
    experience_blocks: list[dict] = []
    education_blocks: list[dict] = []
    achievements: list[str] = []
    contact: dict = {}

    if resume_text and len(resume_text.strip()) > 100:
        try:
            raw = await smart_chat(
                _RESUME_PARSE_PROMPT,
                f"Resume text:\n{resume_text[:5000]}",
                max_tokens=3000, temperature=0.1,
                task_type="resume_parse", cache_ttl=0,
            )
            parsed = _extract_json_safely(raw)
            if isinstance(parsed, dict):
                experience_blocks = parsed.get("experience") or []
                education_blocks = parsed.get("education") or []
                achievements = parsed.get("achievements") or []
                contact = parsed.get("contact") or {}
            elif isinstance(parsed, list):
                # Old format — just an array of experience
                experience_blocks = parsed
        except Exception as e:
            logger.warning(f"Resume parse LLM failed: {e}")

    pdf_bytes = _build_master_pdf(profile, summary, experience_blocks, education_blocks, achievements, contact)
    docx_bytes = _build_master_docx(profile, summary, experience_blocks, education_blocks, achievements, contact)

    return {
        "summary": summary,
        "experience_blocks": experience_blocks,
        "education_blocks": education_blocks,
        "achievements": achievements,
        "pdf_base64": base64.b64encode(pdf_bytes).decode() if pdf_bytes else None,
        "docx_base64": base64.b64encode(docx_bytes).decode() if docx_bytes else None,
    }


async def _analyze_gaps(profile_skills: list, tech_stack: list) -> dict:
    """Use LLM to identify serious gaps between profile and JD."""
    from core.llm import smart_chat
    system_prompt = "You are an ATS Gap Analyst. Compare the profile skills to the job tech stack. Identify 'missing' critical skills and 'matched' skills. Return ONLY JSON: {matched: [skill], missing: [skill], suggestions: [text]}"
    try:
        user_msg = f"Profile: {', '.join(profile_skills)}\nJD Tech: {', '.join(tech_stack)}"
        raw = await smart_chat(system_prompt, user_msg, temperature=0.1, task_type="ats_gap")
        clean = raw.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        return json.loads(clean)
    except Exception:
        return {"matched": [], "missing": tech_stack, "suggestions": "Focus on highlighting core automation skills."}


def _estimate_ats_score(profile: dict, tech_stack: list, gap_analysis: dict = None) -> int:
    """Advanced scoring considering keyword variety and experience overlap."""
    profile_skills = set(
        s.lower() for s in (
            (profile.get("skills") or [])
            + (profile.get("frameworks") or [])
            + (profile.get("cicd_tools") or [])
        )
    )
    # Simple synonym mapping for common tech
    synonyms = {
        "postgres": ["postgresql", "sql", "rdbms"],
        "react": ["frontend", "javascript", "typescript"],
        "aws": ["cloud", "azure", "gcp"],
        "jenkins": ["ci/cd", "github actions", "pipelines"],
    }
    
    jd_techs = set(t.lower() for t in tech_stack)
    if not jd_techs:
        return 75
        
    overlap_count = 0
    for t in jd_techs:
        if t in profile_skills:
            overlap_count += 1
        else:
            # Check synonyms
            for main, alt_list in synonyms.items():
                if t == main and any(alt in profile_skills for alt in alt_list):
                    overlap_count += 0.5  # Partial credit for synonym
                    break

    match_ratio = overlap_count / len(jd_techs)
    base_score = int(match_ratio * 70) + 15  # Base 15, Max 85 from skills
    
    # Bonus for experience
    exp = profile.get("experience_years", 0)
    if exp > 8: base_score += 10
    elif exp > 4: base_score += 5
    
    return min(99, base_score)


class ATSResumeGenerator:
    def __init__(self):
        pass

    async def generate_tailored_resume(self, base_profile: dict, job: dict) -> dict:
        """Uses LLM to rewrite experience bullets to match JD keywords.
        Returns bullets, pdf_base64, docx_base64, and ats_score."""
        import base64
        job_description = job.get("description", "")
        tech_stack = job.get("technologies") or []
        profile_skills = base_profile.get("skills", [])

        # Step 1: LLM Content Generation
        content = await _generate_content(base_profile, job_description, tech_stack)
        bullets = content.get("bullets", [])
        summary = content.get("summary", "")
        skills_grouped = content.get("skills_grouped", {})

        # Step 2: Intelligent Gap Analysis
        gap_analysis = await _analyze_gaps(profile_skills, tech_stack)

        # Step 3: Build Documents
        pdf_bytes = _build_pdf(base_profile, job, bullets, summary=summary, skills_grouped=skills_grouped)
        docx_bytes = _build_docx(base_profile, job, bullets, summary=summary, skills_grouped=skills_grouped)

        return {
            "bullets": bullets,
            "summary": summary,
            "skills_grouped": skills_grouped,
            "pdf_base64": base64.b64encode(pdf_bytes).decode() if pdf_bytes else None,
            "docx_base64": base64.b64encode(docx_bytes).decode() if docx_bytes else None,
            "ats_score": _estimate_ats_score(base_profile, tech_stack, gap_analysis),
            "keyword_heatmap": gap_analysis,
        }
